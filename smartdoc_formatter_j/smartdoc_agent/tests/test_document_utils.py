import unittest
import os
from docx import Document as PythonDocXDocument # Explicit alias
from docx.shared import Pt
# from docx.enum.text import WD_ALIGN_PARAGRAPH # Will be imported below after sys.path manipulation

# Attempt to import utils, handling potential import errors if run standalone
try:
    from smartdoc_agent.utils.document_utils import (
        load_document,
        save_document,
        extract_text_from_paragraphs,
        extract_headings,
        get_document_analysis,
        set_paragraph_font_properties,
        set_paragraph_spacing_properties,
        set_paragraph_alignment_properties,
        apply_set_font_action,
        apply_set_heading_style_action,
        add_table
    )
    from docx.enum.text import WD_ALIGN_PARAGRAPH # Import here if main path works
except ImportError: # pragma: no cover
    # This block allows running tests directly if the package structure isn't fully recognized
    import sys
    import os # Ensure os is imported here if not globally
    sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..'))) # Go up to smartdoc_formatter_j
    from smartdoc_agent.utils.document_utils import (
        load_document,
        save_document,
        extract_text_from_paragraphs,
        extract_headings,
        get_document_analysis,
        set_paragraph_font_properties,
        set_paragraph_spacing_properties,
        set_paragraph_alignment_properties,
        apply_set_font_action,
        apply_set_heading_style_action,
        add_table
    )
    from docx.enum.text import WD_ALIGN_PARAGRAPH # Also import here for the fallback path

class TestDocumentUtils(unittest.TestCase):

    @classmethod
    def setUpClass(cls):
        cls.test_dir = "temp_test_docs_utils"
        os.makedirs(cls.test_dir, exist_ok=True)

        cls.doc_path = os.path.join(cls.test_dir, "test_doc.docx")

        # Create a sample document
        doc = PythonDocXDocument()
        doc.add_heading("Main Title", level=0) # Title style
        doc.add_paragraph("This is the first paragraph.")
        h1 = doc.add_heading("Heading Level 1", level=1)
        h1.runs[0].bold = True # Manually make it bold for later checks if needed
        doc.add_paragraph("Paragraph under H1.")
        doc.add_heading("Another Heading 1", level=1)
        h2 = doc.add_heading("Heading Level 2", level=2)
        h2.runs[0].font.size = Pt(14) # Manually set size
        doc.add_paragraph("Paragraph under H2.")
        doc.add_paragraph("A final paragraph with some text.")

        # Add a paragraph with specific font for more advanced tests later
        # p_font_test = doc.add_paragraph("Font test paragraph.")
        # run = p_font_test.runs[0]
        # run.font.name = 'Arial'
        # run.font.size = Pt(12)

        doc.save(cls.doc_path)

    @classmethod
    def tearDownClass(cls):
        import shutil
        if os.path.exists(cls.test_dir):
            shutil.rmtree(cls.test_dir)

    def test_load_and_save_document(self):
        """Test loading and saving a document."""
        loaded_doc = load_document(self.doc_path)
        self.assertIsNotNone(loaded_doc)
        # self.assertIsInstance(loaded_doc, PythonDocXDocument) # Original line causing error
        self.assertEqual(type(loaded_doc).__name__, 'Document', "loaded_doc should be a Document type")
        self.assertTrue(type(loaded_doc).__module__.startswith('docx.document'), "loaded_doc module should be from docx.document")

        save_path = os.path.join(self.test_dir, "saved_test_doc.docx")
        save_document(loaded_doc, save_path)
        self.assertTrue(os.path.exists(save_path))

        # Try loading the saved document to ensure it's valid
        reloaded_doc = load_document(save_path)
        self.assertIsNotNone(reloaded_doc)
        self.assertEqual(type(reloaded_doc).__name__, 'Document', "reloaded_doc should be a Document type")
        self.assertTrue(type(reloaded_doc).__module__.startswith('docx.document'), "reloaded_doc module should be from docx.document")

    def test_extract_text_from_paragraphs(self):
        """Test extracting text from paragraphs."""
        doc = load_document(self.doc_path)
        texts = extract_text_from_paragraphs(doc)

        expected_texts = [
            "Main Title", # Headings are also paragraphs in python-docx
            "This is the first paragraph.",
            "Heading Level 1",
            "Paragraph under H1.",
            "Another Heading 1",
            "Heading Level 2",
            "Paragraph under H2.",
            "A final paragraph with some text."
        ]
        self.assertEqual(texts, expected_texts)

    def test_extract_headings(self):
        """Test extracting headings and their levels."""
        doc = load_document(self.doc_path)
        headings = extract_headings(doc)

        # Note: "Main Title" (level 0) might be "Title" or "Heading 0" style
        # The current extract_headings only looks for "Heading X"
        # Let's adjust expectations or the function if level 0 is critical.
        # For now, assuming default "Heading X" styles are primary targets.

        expected_headings = [
            {"level": 1, "text": "Heading Level 1"},
            {"level": 1, "text": "Another Heading 1"},
            {"level": 2, "text": "Heading Level 2"}
        ]
        # Filter out level 0 if it's picked up, or adjust based on actual style name
        # For now, the function `para.style.name.startswith("Heading")` is key

        # If "Main Title" uses a style like "Title" it won't be caught.
        # If it uses "Heading 0" (some templates might), it would.
        # Let's check the created document's style for "Main Title"
        # default doc.add_heading(level=0) -> "Title" style
        # default doc.add_heading(level=1) -> "Heading 1" style

        # So, the current extract_headings will NOT pick up the "Main Title" (level 0)
        # The new get_document_analysis does, so this old test for extract_headings might need adjustment
        # or removal if extract_headings is no longer used directly.
        # For now, let's assume extract_headings is still a standalone utility for some cases.
        doc_for_headings_test = PythonDocXDocument()
        doc_for_headings_test.add_heading("H1", level=1)
        doc_for_headings_test.add_heading("H2", level=2)
        actual_headings = extract_headings(doc_for_headings_test)
        expected_headings_standalone = [
            {"level": 1, "text": "H1"},
            {"level": 2, "text": "H2"}
        ]
        self.assertEqual(actual_headings, expected_headings_standalone)


    def test_get_document_analysis(self):
        """Test the full get_document_analysis function."""
        doc = load_document(self.doc_path)
        analysis = get_document_analysis(doc)

        self.assertIn("elements", analysis)
        elements = analysis["elements"]
        self.assertTrue(len(elements) > 0)

        # Check structure of the first element (should be "Main Title" - level 0 heading)
        first_element = elements[0]
        self.assertEqual(first_element["type"], "heading")
        self.assertEqual(first_element["level"], 0)
        self.assertEqual(first_element["text"], "Main Title")
        self.assertIn("style_name", first_element) # e.g., 'Title'
        self.assertTrue(len(first_element["runs"]) > 0)

        # Check a standard paragraph
        # Paragraphs are: "Main Title", "This is the first paragraph.", "Heading Level 1", ...
        # So elements[1] should be "This is the first paragraph."
        para_element = elements[1]
        self.assertEqual(para_element["type"], "paragraph")
        self.assertEqual(para_element["text"], "This is the first paragraph.")
        self.assertIn("style_name", para_element)
        self.assertTrue(len(para_element["runs"]) > 0)
        first_run_details = para_element["runs"][0]
        self.assertIn("font_name", first_run_details)
        self.assertIn("font_size", first_run_details)
        self.assertIn("bold", first_run_details)

        # Check a styled heading (Heading Level 1)
        h1_element = elements[2] # "Main Title", "Para1", "Heading Level 1"
        self.assertEqual(h1_element["type"], "heading")
        self.assertEqual(h1_element["level"], 1)
        self.assertEqual(h1_element["text"], "Heading Level 1")
        # self.assertTrue(h1_element["runs"][0]["bold"]) # This was manually set in setUpClass, useful for detailed run checks

    def test_load_document_file_not_found(self):
        """Test loading a non-existent document."""
        with self.assertRaises(Exception): # python-docx raises PackageNotFoundError or similar, Exception is broad
            load_document("non_existent_document_for_sure.docx")

    def test_set_paragraph_font_properties(self):
        doc = PythonDocXDocument()
        p = doc.add_paragraph("Test font change.")
        set_paragraph_font_properties(p, font_name="Arial", size_pt=12, bold=True)

        self.assertTrue(len(p.runs) > 0, "Paragraph should have runs after adding text.")
        # Check properties of the first run (assuming text isn't further broken down)
        run = p.runs[0]
        self.assertEqual(run.font.name, "Arial")
        self.assertEqual(run.font.size, Pt(12))
        self.assertTrue(run.bold)
        self.assertIsNone(run.italic) # Check not inadvertently set

    def test_set_paragraph_spacing_properties(self):
        doc = PythonDocXDocument()
        p = doc.add_paragraph("Test spacing change.")
        set_paragraph_spacing_properties(p, spacing_before_pt=6, spacing_after_pt=12, line_spacing_rule=1.5)

        pf = p.paragraph_format
        self.assertEqual(pf.space_before, Pt(6))
        self.assertEqual(pf.space_after, Pt(12))
        self.assertEqual(pf.line_spacing, 1.5)

    def test_set_paragraph_alignment_properties(self):
        doc = PythonDocXDocument()
        p = doc.add_paragraph("Test alignment change.")
        set_paragraph_alignment_properties(p, alignment="CENTER")
        self.assertEqual(p.alignment, WD_ALIGN_PARAGRAPH.CENTER)

        set_paragraph_alignment_properties(p, alignment="INVALID_ALIGN") # Test invalid value
        self.assertEqual(p.alignment, WD_ALIGN_PARAGRAPH.CENTER) # Should remain unchanged

    # Tests for higher-level apply_..._action functions (these are more integration-like for utils)
    def test_apply_set_font_action(self):
        doc = load_document(self.doc_path) # Uses the doc created in setUpClass
        analysis = get_document_analysis(doc) # Get analysis to simulate tool environment

        # Action to change font of all body paragraphs (not headings)
        action = {"action": "set_font", "scope": "all_body_paragraphs", "font_name": "Courier New", "size": 10, "italic": True}
        apply_set_font_action(doc, analysis["elements"], action)

        # Verification (simplified: check one known body paragraph)
        # elements[1] is "This is the first paragraph."
        first_body_para_idx = analysis["elements"][1]["paragraph_index"]
        first_body_para = doc.paragraphs[first_body_para_idx]
        self.assertEqual(first_body_para.runs[0].font.name, "Courier New")
        self.assertEqual(first_body_para.runs[0].font.size, Pt(10))
        self.assertTrue(first_body_para.runs[0].italic)

        # Check that a heading was NOT changed by "all_body_paragraphs"
        # elements[0] is "Main Title" (heading)
        first_heading_idx = analysis["elements"][0]["paragraph_index"]
        first_heading_para = doc.paragraphs[first_heading_idx]
        self.assertNotEqual(first_heading_para.runs[0].font.name, "Courier New")


    def test_apply_set_heading_style_action(self):
        doc = load_document(self.doc_path)
        analysis = get_document_analysis(doc)

        action = {"action": "set_heading_style", "level": 1, "font_name": "Impact", "size": 20, "bold": False, "spacing_after": 18}
        apply_set_heading_style_action(doc, analysis["elements"], action)

        # Verification: Check a known H1
        # elements[2] is "Heading Level 1"
        h1_idx = analysis["elements"][2]["paragraph_index"]
        h1_para = doc.paragraphs[h1_idx]
        self.assertEqual(h1_para.runs[0].font.name, "Impact")
        self.assertEqual(h1_para.runs[0].font.size, Pt(20))
        self.assertFalse(h1_para.runs[0].bold) # Assuming it might have been bold before
        self.assertEqual(h1_para.paragraph_format.space_after, Pt(18))

    def test_add_table_utility(self):
        """Test the add_table utility function."""
        doc = PythonDocXDocument()

        # Test basic table creation
        table1 = add_table(doc, rows=2, cols=3)
        self.assertIsNotNone(table1)
        self.assertEqual(len(table1.rows), 2)
        self.assertEqual(len(table1.columns), 3)
        self.assertEqual(len(doc.tables), 1)

        # Test table creation with data
        data = [["Name", "Age"], ["Alice", "30"], ["Bob", "24"]]
        table2 = add_table(doc, rows=3, cols=2, data=data)
        self.assertIsNotNone(table2)
        self.assertEqual(len(table2.rows), 3)
        self.assertEqual(len(table2.columns), 2)
        self.assertEqual(table2.cell(0, 0).text, "Name")
        self.assertEqual(table2.cell(1, 1).text, "30")
        self.assertEqual(table2.cell(2, 0).text, "Bob")
        self.assertEqual(len(doc.tables), 2)

        # Test data dimensions mismatch (more rows in data than table rows)
        data_more_rows = [["R1C1"], ["R2C1"], ["R3C1"]]
        table_data_mr = add_table(doc, rows=2, cols=1, data=data_more_rows)
        self.assertEqual(table_data_mr.cell(0,0).text, "R1C1")
        self.assertEqual(table_data_mr.cell(1,0).text, "R2C1")
        # R3C1 should be ignored

        # Test data dimensions mismatch (more cols in data than table cols)
        data_more_cols = [["R1C1", "R1C2", "R1C3"]]
        table_data_mc = add_table(doc, rows=1, cols=2, data=data_more_cols)
        self.assertEqual(table_data_mc.cell(0,0).text, "R1C1")
        self.assertEqual(table_data_mc.cell(0,1).text, "R1C2")
        # R1C3 should be ignored

        # Test table creation with a valid style
        # Note: For this to pass robustly, the style name must be valid in typical Word installations.
        # 'Table Grid' is a very common and safe one.
        table3 = add_table(doc, rows=2, cols=2, style='Table Grid')
        self.assertIsNotNone(table3)
        self.assertEqual(table3.style.name, 'Table Grid') # python-docx might store it as 'TableGrid' or similar

        # Test with an invalid style (python-docx usually defaults or ignores, doesn't hard error)
        # The `add_table` utility function prints a warning.
        table4 = add_table(doc, rows=1, cols=1, style='NonExistentStyle123')
        self.assertIsNotNone(table4) # Table should still be created
        # Default style might be applied, or style might be None or have a default name
        # Exact assertion depends on python-docx behavior for truly invalid styles.
        # For now, ensuring it doesn't crash and returns a table is sufficient.
        # The utility function itself prints a warning which is good.

        # Test invalid rows/cols
        invalid_table_rows = add_table(doc, rows=0, cols=2)
        self.assertIsNone(invalid_table_rows)
        invalid_table_cols = add_table(doc, rows=2, cols=-1)
        self.assertIsNone(invalid_table_cols)

        # Test invalid document object
        not_a_doc = "I am not a document"
        invalid_doc_call = add_table(not_a_doc, rows=2, cols=2)
        self.assertIsNone(invalid_doc_call)


if __name__ == '__main__':
    unittest.main()
