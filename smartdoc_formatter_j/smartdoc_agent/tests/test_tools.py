import unittest
import os
import json # Added to fix NameError
from unittest.mock import patch, MagicMock
from docx import Document as PythonDocXDocument

# Attempt to import tools, handling potential import errors if run standalone
try:
    from smartdoc_agent.core.tools import (
        analyze_document_structure,
        create_formatting_plan,
        apply_contextual_formatting,
        validate_formatting_result,
        create_table,
        format_table_cell_tool,
        merge_table_cells_tool
    )
    # document_utils are mocked, but good to have them if direct calls were made
    from smartdoc_agent.utils.document_utils import (
        load_document,
        save_document,
        add_table,
        format_table_cell,
        merge_table_cells
    )
except ImportError: # pragma: no cover
    # This is to allow running the test file directly for development,
    # assuming it's in the tests/ directory and smartdoc_agent is in the parent.
    import sys
    sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..'))) # Go up to smartdoc_formatter_j
    from smartdoc_agent.core.tools import (
        analyze_document_structure,
        create_formatting_plan,
        apply_contextual_formatting,
        validate_formatting_result,
        create_table,
        format_table_cell_tool,
        merge_table_cells_tool
    )
    from smartdoc_agent.utils.document_utils import (
        load_document,
        save_document,
        add_table,
        format_table_cell,
        merge_table_cells
    )


class TestFormattingTools(unittest.TestCase):

    @classmethod
    def setUpClass(cls):
        """Set up for all tests; create a sample document."""
        cls.test_doc_path = "temp_test_document_for_tools.docx"
        cls.test_output_doc_path = "temp_test_document_for_tools_modified.docx"

        doc = PythonDocXDocument()
        doc.add_heading("Test Document Title", level=1)
        doc.add_paragraph("This is a paragraph for testing.")
        doc.add_heading("Another Heading", level=2)
        doc.add_paragraph("Second paragraph with some more text.")
        save_document(doc, cls.test_doc_path)

    @classmethod
    def tearDownClass(cls):
        """Clean up after all tests; remove the sample document."""
        if os.path.exists(cls.test_doc_path):
            os.remove(cls.test_doc_path)
        if os.path.exists(cls.test_output_doc_path):
            os.remove(cls.test_output_doc_path)

    def test_analyze_document_structure_runs(self):
        """Test that analyze_document_structure runs and returns a JSON string."""
        import json
        # Langchain tools expect the input dict to be under a key matching the Pydantic model field,
        # which is often 'tool_input' if the function takes a single dict/str argument.
        tool_args = {"doc_path": self.test_doc_path}
        result_json_str = analyze_document_structure.invoke({"tool_input": tool_args})
        self.assertIsInstance(result_json_str, str)

        try:
            result_data = json.loads(result_json_str)
        except json.JSONDecodeError:
            self.fail("analyze_document_structure did not return valid JSON.")

        self.assertIsInstance(result_data, dict)
        self.assertEqual(result_data.get("document_path"), self.test_doc_path)
        self.assertIn("elements", result_data)
        self.assertIn("summary", result_data)
        self.assertTrue(len(result_data["elements"]) > 0)
        self.assertTrue(result_data["summary"]["total_elements"] > 0)
        # Check if one of the known texts from setUpClass is present in the analysis
        found_title = any(el["text"] == "Test Document Title" for el in result_data["elements"])
        self.assertTrue(found_title, "Expected content not found in analysis elements.")


    @patch('smartdoc_agent.core.tools.llm') # Mock the LLM used by the tool
    def test_create_formatting_plan_runs(self, mock_llm):
        """Test that create_formatting_plan runs and returns a JSON string plan."""
        import json

        # Mock the LLM's response
        mock_llm_response_content = json.dumps([
            {"action": "set_font", "scope": "all_body_paragraphs", "font_name": "Arial", "size": 12},
            {"action": "set_heading_style", "level": 1, "font_name": "Arial Black", "size": 16, "bold": True}
        ])
        mock_llm.invoke.return_value = MagicMock(content=mock_llm_response_content)

        # Sample analysis (must be a JSON string)
        sample_analysis_dict = {
            "document_path": self.test_doc_path,
            "summary": {"total_elements": 4, "paragraph_count": 2, "heading_count": 2},
            "elements": [
                {"type": "heading", "level": 1, "text": "Test Document Title", "style_name": "Heading 1", "runs": []},
                {"type": "paragraph", "text": "This is a paragraph for testing.", "style_name": "Normal", "runs": []}
            ]
        }
        sample_analysis_json = json.dumps(sample_analysis_dict)
        user_goal = "Make it look professional."

        tool_args = {
            "document_analysis_json": sample_analysis_json,
            "user_goal": user_goal
        }
        result_plan_json = create_formatting_plan.invoke({"tool_input": tool_args})

        self.assertIsInstance(result_plan_json, str)
        try:
            plan_data = json.loads(result_plan_json)
        except json.JSONDecodeError:
            self.fail("create_formatting_plan did not return valid JSON.")

        self.assertIsInstance(plan_data, list)
        self.assertTrue(len(plan_data) > 0)
        self.assertIn("action", plan_data[0])
        mock_llm.invoke.assert_called_once() # Check that LLM was called

    def test_apply_contextual_formatting_runs(self):
        """Test that apply_contextual_formatting runs and returns a string."""
        import json
        # This is a basic test for the stub.
        # The plan should now be a JSON string.
        sample_plan_list = [
            {"action": "set_font", "scope": "all_body_paragraphs", "font_name": "Times New Roman", "size": 12},
            {"action": "set_heading_style", "level": 1, "font_name": "TNR Bold", "size": 14, "bold": True}
        ]
        sample_plan_json = json.dumps(sample_plan_list)

        # Sample analysis JSON string (needed by the refined tool)
        sample_analysis_dict = {
            "document_path": self.test_doc_path,
            "summary": {"total_elements": 1, "paragraph_count": 1, "heading_count": 0},
            "elements": [{"type": "paragraph", "text": "Text", "paragraph_index": 0}]
        }
        sample_analysis_json = json.dumps(sample_analysis_dict)

        # Define a side effect function for the save_document mock
        def mock_save_document_side_effect(doc, path):
            # Create a dummy file to make os.path.exists() return True
            with open(path, 'w') as f:
                f.write("dummy content")
            return True # Simulate successful save

        # Mock the document utility functions that perform changes
        with patch('smartdoc_agent.core.tools.load_document') as mock_load, \
             patch('smartdoc_agent.core.tools.save_document', side_effect=mock_save_document_side_effect) as mock_save, \
             patch('smartdoc_agent.core.tools.apply_set_font_action') as mock_apply_font, \
             patch('smartdoc_agent.core.tools.apply_set_heading_style_action') as mock_apply_heading:

            mock_doc_instance = MagicMock(spec=PythonDocXDocument)
            mock_load.return_value = mock_doc_instance

            tool_args = {
                "doc_path": self.test_doc_path,
                "formatting_plan_json": sample_plan_json,
                "document_analysis_json": sample_analysis_json,
                "output_doc_path": self.test_output_doc_path
            }
            result_json_str = apply_contextual_formatting.invoke({"tool_input": tool_args})
            self.assertIsInstance(result_json_str, str)

            try:
                result_data = json.loads(result_json_str)
            except json.JSONDecodeError:
                self.fail("apply_contextual_formatting did not return valid JSON.")

            self.assertEqual(result_data.get("status"), "success")
            # This assertion now passes because our mock_save_document_side_effect creates the file
            self.assertTrue(os.path.exists(self.test_output_doc_path))

            # Check if our action handlers were called
            mock_load.assert_called_once_with(self.test_doc_path)
            mock_apply_font.assert_called_once()
            mock_apply_heading.assert_called_once()
            mock_save.assert_called_once_with(mock_doc_instance, self.test_output_doc_path)


    def test_validate_formatting_result_runs(self):
        """Test that validate_formatting_result runs and returns a string."""
        # This is a basic test for the stub.
        # Ensure the modified document exists for validation
        if not os.path.exists(self.test_output_doc_path):
             # Create a dummy one if apply_contextual_formatting hasn't created it yet
            doc = load_document(self.test_doc_path)
            save_document(doc, self.test_output_doc_path)

        sample_plan = "1. Change font. 2. Check headings." # This tool's old stub took different args
        user_goal = "Check professionalism."

        # The actual validate_formatting_result tool expects JSON strings for analyses and plan
        dummy_analysis_json = json.dumps({"document_path": "dummy.docx", "summary": {}, "elements": []})
        dummy_plan_json = json.dumps([{"action":"dummy"}])

        tool_args = {
            "original_doc_analysis_json": dummy_analysis_json,
            "modified_doc_analysis_json": dummy_analysis_json,
            "formatting_plan_json": dummy_plan_json,
            "user_goal": user_goal
        }
        # Mock the LLM for validate_formatting_result as it uses one
        with patch('smartdoc_agent.core.tools.llm') as mock_validate_llm:
            mock_validate_llm.invoke.return_value = MagicMock(content=json.dumps({"overall_assessment": "Good"}))
            result = validate_formatting_result.invoke({"tool_input": tool_args})

        self.assertIsInstance(result, str)
        # The tool now returns raw JSON from the LLM, so we check the JSON content.
        result_data = json.loads(result)
        self.assertEqual(result_data.get("overall_assessment"), "Good")


class TestCreateTableTool(unittest.TestCase):
    def setUp(self):
        self.test_doc_path = "temp_create_table_input.docx"
        self.test_output_doc_path = "temp_create_table_output.docx"

        # Create a dummy input document for tests that need one
        doc = PythonDocXDocument()
        doc.add_paragraph("Initial content.")
        doc.save(self.test_doc_path)

    def tearDown(self):
        if os.path.exists(self.test_doc_path):
            os.remove(self.test_doc_path)
        if os.path.exists(self.test_output_doc_path):
            os.remove(self.test_output_doc_path)

    @patch('smartdoc_agent.core.tools.save_document')
    @patch('smartdoc_agent.core.tools.add_table')
    @patch('smartdoc_agent.core.tools.load_document')
    def test_create_table_success_basic(self, mock_load_document, mock_add_table, mock_save_document):
        import json
        mock_doc = MagicMock(spec=PythonDocXDocument)
        mock_load_document.return_value = mock_doc

        mock_table_obj = MagicMock()
        mock_table_obj.rows = 2 # Simulate table object attributes
        mock_table_obj.columns = 3
        mock_add_table.return_value = mock_table_obj

        tool_input = {
            "doc_path": self.test_doc_path,
            "rows": 2,
            "cols": 3,
            "output_doc_path": self.test_output_doc_path
        }
        result_json = create_table.invoke({"tool_input": tool_input})
        result = json.loads(result_json)

        self.assertEqual(result["status"], "success")
        self.assertIn("Table with 2x3 created", result["message"])
        self.assertEqual(result["output_doc_path"], self.test_output_doc_path)
        self.assertIsNotNone(result["table_details"])
        self.assertEqual(result["table_details"]["rows"], 2)
        self.assertEqual(result["table_details"]["cols"], 3)
        self.assertEqual(result["table_details"]["style_applied"], "Default")

        mock_load_document.assert_called_once_with(self.test_doc_path)
        mock_add_table.assert_called_once_with(mock_doc, 2, 3, data=None, style=None)
        mock_save_document.assert_called_once_with(mock_doc, self.test_output_doc_path)

    @patch('smartdoc_agent.core.tools.save_document')
    @patch('smartdoc_agent.core.tools.add_table')
    @patch('smartdoc_agent.core.tools.load_document')
    def test_create_table_success_with_data_and_style(self, mock_load_document, mock_add_table, mock_save_document):
        import json
        mock_doc = MagicMock(spec=PythonDocXDocument)
        mock_load_document.return_value = mock_doc

        mock_table_obj = MagicMock()
        mock_table_obj.rows = 1
        mock_table_obj.columns = 1
        mock_add_table.return_value = mock_table_obj

        table_data = [["Test Data"]]
        table_style = "Table Grid"

        tool_input = {
            "doc_path": self.test_doc_path,
            "rows": 1,
            "cols": 1,
            "data": table_data,
            "style": table_style,
            "output_doc_path": self.test_output_doc_path
        }
        result_json = create_table.invoke({"tool_input": tool_input})
        result = json.loads(result_json)

        self.assertEqual(result["status"], "success")
        self.assertEqual(result["table_details"]["style_applied"], table_style)
        mock_add_table.assert_called_once_with(mock_doc, 1, 1, data=table_data, style=table_style)

    def test_create_table_invalid_input_args(self):
        import json
        # Missing rows
        tool_input_missing_rows = {
            "doc_path": self.test_doc_path, "cols": 2, "output_doc_path": self.test_output_doc_path
        }
        result_json = create_table.invoke({"tool_input": tool_input_missing_rows})
        result = json.loads(result_json)
        self.assertEqual(result["status"], "error")
        self.assertIn("Missing or invalid required arguments", result["message"])

        # Invalid rows (not positive)
        tool_input_invalid_rows = {
            "doc_path": self.test_doc_path, "rows": 0, "cols": 2, "output_doc_path": self.test_output_doc_path
        }
        result_json = create_table.invoke({"tool_input": tool_input_invalid_rows})
        result = json.loads(result_json)
        self.assertEqual(result["status"], "error")
        self.assertIn("'rows' (0) and 'cols' (2) must be positive integers", result["message"])

    @patch('smartdoc_agent.core.tools.load_document')
    def test_create_table_load_failure(self, mock_load_document):
        import json
        mock_load_document.side_effect = FileNotFoundError("Mocked File Not Found")
        tool_input = {"doc_path": "nonexistent.docx", "rows": 2, "cols": 2, "output_doc_path": self.test_output_doc_path}
        result_json = create_table.invoke({"tool_input": tool_input})
        result = json.loads(result_json)
        self.assertEqual(result["status"], "error")
        self.assertIn("Input document not found at nonexistent.docx", result["message"])

    @patch('smartdoc_agent.core.tools.load_document')
    @patch('smartdoc_agent.core.tools.add_table')
    def test_create_table_add_table_failure(self, mock_add_table, mock_load_document):
        import json
        mock_doc = MagicMock(spec=PythonDocXDocument)
        mock_load_document.return_value = mock_doc
        mock_add_table.return_value = None # Simulate failure in add_table utility

        tool_input = {"doc_path": self.test_doc_path, "rows": 2, "cols": 2, "output_doc_path": self.test_output_doc_path}
        result_json = create_table.invoke({"tool_input": tool_input})
        result = json.loads(result_json)
        self.assertEqual(result["status"], "error")
        self.assertIn("Failed to create table using document_utils.add_table", result["message"])

class TestTableManipulationTools(unittest.TestCase):
    def setUp(self):
        self.doc_path = "temp_table_tools_doc.docx"
        self.output_path = "temp_table_tools_doc_modified.docx"
        doc = PythonDocXDocument()
        doc.add_table(rows=2, cols=2) # Add a table for the tools to find
        doc.save(self.doc_path)

    def tearDown(self):
        if os.path.exists(self.doc_path):
            os.remove(self.doc_path)
        if os.path.exists(self.output_path):
            os.remove(self.output_path)

    @patch('smartdoc_agent.core.tools.save_document')
    @patch('smartdoc_agent.core.tools.format_table_cell')
    @patch('smartdoc_agent.core.tools.load_document')
    def test_format_table_cell_tool_success(self, mock_load, mock_format, mock_save):
        mock_doc = MagicMock()
        mock_table = MagicMock()
        mock_doc.tables = [mock_table]
        mock_load.return_value = mock_doc
        mock_format.return_value = True

        tool_args = {
            "doc_path": self.doc_path, "output_doc_path": self.output_path,
            "table_index": 0, "row": 0, "col": 0, "text": "Hello"
        }
        result = json.loads(format_table_cell_tool.invoke({"tool_input": tool_args}))
        self.assertEqual(result["status"], "success")
        mock_format.assert_called_once_with(table=mock_table, row=0, col=0, text="Hello", font_name=None, font_size=None, bold=None, italic=None, underline=None, alignment=None, shading=None)
        mock_save.assert_called_once()

    def test_format_table_cell_tool_invalid_index(self):
        tool_args = {
            "doc_path": self.doc_path, "output_doc_path": self.output_path,
            "table_index": 99, "row": 0, "col": 0, "text": "Hello"
        }
        result = json.loads(format_table_cell_tool.invoke({"tool_input": tool_args}))
        self.assertEqual(result["status"], "error")
        self.assertIn("Table index 99 is out of bounds", result["message"])

    @patch('smartdoc_agent.core.tools.save_document')
    @patch('smartdoc_agent.core.tools.merge_table_cells')
    @patch('smartdoc_agent.core.tools.load_document')
    def test_merge_table_cells_tool_success(self, mock_load, mock_merge, mock_save):
        mock_doc = MagicMock()
        mock_table = MagicMock()
        mock_doc.tables = [mock_table]
        mock_load.return_value = mock_doc
        mock_merge.return_value = True

        tool_args = {
            "doc_path": self.doc_path, "output_doc_path": self.output_path,
            "table_index": 0, "start_row": 0, "start_col": 0, "end_row": 0, "end_col": 1
        }
        result = json.loads(merge_table_cells_tool.invoke({"tool_input": tool_args}))
        self.assertEqual(result["status"], "success")
        mock_merge.assert_called_once_with(mock_table, 0, 0, 0, 1)
        mock_save.assert_called_once()


if __name__ == '__main__':
    unittest.main()
