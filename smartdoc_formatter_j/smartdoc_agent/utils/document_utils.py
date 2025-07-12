import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.table import Table, _Cell
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- Document Load/Save Utilities ---

def load_document(file_path: str) -> Document:
    """Loads a .docx document from the given file path."""
    try:
        document = Document(file_path)
        return document
    except Exception as e:
        print(f"Error loading document {file_path}: {e}")
        raise

def save_document(document: Document, file_path: str) -> None:
    """Saves the given Document object to a .docx file."""
    try:
        document.save(file_path)
        print(f"Document saved to {file_path}")
    except Exception as e:
        print(f"Error saving document to {file_path}: {e}")
        raise

def extract_text_from_paragraphs(document: Document) -> list[str]:
    """Extracts text from all paragraphs in the document."""
    texts = []
    for para in document.paragraphs:
        texts.append(para.text)
    return texts

def extract_headings(document: Document) -> list[dict[str, str]]:
    """
    Extracts headings and their levels from the document.
    Assumes headings are styled with Word's default heading styles (e.g., "Heading 1", "Heading 2").
    """
    headings = []
    for para in document.paragraphs:
        if para.style and para.style.name.startswith("Heading"):
            try:
                level = int(para.style.name.split(" ")[-1])
            except ValueError:
                level = 0 # Or handle as an unknown heading level
            headings.append({"level": level, "text": para.text})
    return headings

def get_run_details(run) -> dict:
    """Extracts details from a run."""
    return {
        "text": run.text,
        "font_name": run.font.name,
        "font_size": run.font.size.pt if run.font.size else None, # Size in points
        "bold": run.bold,
        "italic": run.italic,
        "underline": run.underline,
    }

def get_paragraph_details(para, para_index: int) -> dict:
    """Extracts details from a paragraph, including its runs."""
    alignment_name = None
    if para.alignment is not None:
        # para.alignment should be a member of WD_ALIGN_PARAGRAPH enum
        # Accessing .name gives the string representation like 'LEFT', 'CENTER'
        try:
            alignment_name = para.alignment.name
        except AttributeError: # Should not happen if para.alignment is a valid enum member
            alignment_name = str(para.alignment) # Fallback to string of the value

    para_info = {
        "paragraph_index": para_index,
        "text": para.text,
        "style_name": para.style.name if para.style else "Default Paragraph Font",
        "alignment": alignment_name,
        "runs": [get_run_details(run) for run in para.runs]
    }
    # Heading level detection
    if para.style and para.style.name.startswith('Heading'):
        try:
            level = int(para.style.name.split(' ')[-1])
            para_info["type"] = "heading"
            para_info["level"] = level
        except ValueError:
            para_info["type"] = "paragraph" # Could be a custom heading style not ending in a number
            para_info["level"] = 0
    elif para.style and para.style.name == "Title": # Common style for document title
        para_info["type"] = "heading"
        para_info["level"] = 0
    else:
        para_info["type"] = "paragraph"

    return para_info

def get_document_analysis(document: Document) -> dict:
    """
    Analyzes a Document object and extracts detailed information about its elements.
    Returns a dictionary with a list of element details.
    """
    analysis = {"elements": []}
    for i, para in enumerate(document.paragraphs):
        para_details = get_paragraph_details(para, i)
        analysis["elements"].append(para_details)

    # Future: Add analysis for tables, lists, images, sections, etc.
    return analysis

if __name__ == '__main__':
    # This is for basic testing of the utility functions.
    # You would need a sample .docx file named 'sample.docx' in the same directory.
    # Create a dummy document for testing if 'sample.docx' doesn't exist
    doc = Document()
    doc.add_heading('Test Document', level=0)
    doc.add_paragraph('This is a test paragraph.')
    doc.add_heading('Heading 1', level=1)
    doc.add_paragraph('Another paragraph under Heading 1.')
    doc.add_heading('Heading 2', level=2)
    doc.add_paragraph('Paragraph under Heading 2.')
    save_document(doc, 'sample_test_doc.docx')

    print("Testing with 'sample_test_doc.docx'")
    loaded_doc = load_document('sample_test_doc.docx')

    print("\nExtracted Paragraph Texts:")
    texts = extract_text_from_paragraphs(loaded_doc)
    for text in texts:
        print(f"- {text}")

    print("\nExtracted Headings:")
    headings_info = extract_headings(loaded_doc)
    for heading in headings_info:
        print(f"- Level {heading['level']}: {heading['text']}")

    print("\nDetailed Document Analysis:")
    detailed_analysis_data = get_document_analysis(loaded_doc) # Use the loaded_doc directly
    for item in detailed_analysis_data["elements"]:
        print(f"- Type: {item['type']}")
        if item['type'] == 'heading':
            print(f"  Level: {item['level']}")
        print(f"  Text: {item['text'][:50]}...") # Print first 50 chars
        if 'style' in item:
            print(f"  Style: {item['style']}")
        if 'runs' in item:
            # print(f"  Runs: {item['runs']}") # This can be verbose
            if item['runs']:
                first_run = item['runs'][0]
                print(f"  First run font: {first_run.get('font_name', 'N/A')}, Size: {first_run.get('font_size', 'N/A')}, Bold: {first_run.get('bold', 'N/A')}")

    # This is for basic testing of the utility functions.
    # You would need a sample .docx file named 'sample.docx' in the same directory.
    # Create a dummy document for testing if 'sample.docx' doesn't exist
    doc_main = Document() # Renamed to avoid conflict with 'doc' in add_table test
    doc_main.add_heading('Test Document', level=0)
    doc_main.add_paragraph('This is a test paragraph.')
    doc_main.add_heading('Heading 1', level=1)
    doc_main.add_paragraph('Another paragraph under Heading 1.')
    doc_main.add_heading('Heading 2', level=2)
    doc_main.add_paragraph('Paragraph under Heading 2.')
    save_document(doc_main, 'sample_test_doc.docx')

    print("Testing with 'sample_test_doc.docx'")
    loaded_doc = load_document('sample_test_doc.docx')

    print("\nExtracted Paragraph Texts:")
    texts = extract_text_from_paragraphs(loaded_doc)
    for text in texts:
        print(f"- {text}")

    print("\nExtracted Headings:")
    headings_info = extract_headings(loaded_doc)
    for heading in headings_info:
        print(f"- Level {heading['level']}: {heading['text']}")

    print("\nDetailed Document Analysis:")
    detailed_analysis_data = get_document_analysis(loaded_doc) # Use the loaded_doc directly
    for item in detailed_analysis_data["elements"]:
        print(f"- Type: {item['type']}")
        if item['type'] == 'heading':
            print(f"  Level: {item['level']}")
        print(f"  Text: {item['text'][:50]}...") # Print first 50 chars
        if 'style_name' in item: # Corrected from 'style' to 'style_name' for consistency
            print(f"  Style: {item['style_name']}")
        if 'runs' in item:
            if item['runs']:
                first_run = item['runs'][0]
                print(f"  First run font: {first_run.get('font_name', 'N/A')}, Size: {first_run.get('font_size', 'N/A')}, Bold: {first_run.get('bold', 'N/A')}")

    # Test add_table
    print("\n--- Testing add_table utility ---")
    doc_for_table_test = Document()
    doc_for_table_test.add_paragraph("Paragraph before table.")
    table_data = [["Name", "Role"], ["Jules", "Agent"], ["User", "Supervisor"]]
    created_table = add_table(doc_for_table_test, rows=3, cols=2, data=table_data, style="Table Grid")

    if created_table:
        print(f"Test table created with style: {created_table.style.name if created_table.style else 'Default'}")
        # Basic check of table content
        if len(created_table.rows) == 3 and len(created_table.columns) == 2:
            print(f"  Cell (0,0): {created_table.cell(0,0).text}")
            print(f"  Cell (1,1): {created_table.cell(1,1).text}")
        else:
            print("  Error: Table dimensions mismatch after creation.")
    else:
        print("Test table creation FAILED.")
    doc_for_table_test.add_paragraph("Paragraph after table.")

    save_document(doc_for_table_test, 'sample_test_doc_with_table.docx')
    print("Saved 'sample_test_doc_with_table.docx' for manual inspection.")

    # Clean up the dummy files
    import os
    if os.path.exists('sample_test_doc.docx'):
        os.remove('sample_test_doc.docx')
        print("\nCleaned up sample_test_doc.docx")
    if os.path.exists('sample_test_doc_with_table.docx'):
        os.remove('sample_test_doc_with_table.docx')
        print("Cleaned up sample_test_doc_with_table.docx")

# --- Element Analysis Utilities ---
# (This comment indicates the end of the __main__ block and start of next section)

# --- Low-Level Formatting Utilities ---

def set_paragraph_font_properties(paragraph, font_name: str = None, size_pt: float = None, bold: bool = None, italic: bool = None, underline: bool = None):
    """Applies font properties to all runs in a paragraph."""
    for run in paragraph.runs:
        if font_name:
            run.font.name = font_name
        if size_pt:
            run.font.size = Pt(size_pt)
        if bold is not None:
            run.bold = bold
        if italic is not None:
            run.italic = italic
        if underline is not None:
            run.underline = underline

def set_paragraph_spacing_properties(paragraph, spacing_before_pt: float = None, spacing_after_pt: float = None, line_spacing_rule: float = None): # line_spacing_rule e.g. 1.0, 1.5, 2.0
    """Applies spacing properties to a paragraph."""
    if spacing_before_pt is not None:
        paragraph.paragraph_format.space_before = Pt(spacing_before_pt)
    if spacing_after_pt is not None:
        paragraph.paragraph_format.space_after = Pt(spacing_after_pt)
    if line_spacing_rule is not None:
        paragraph.paragraph_format.line_spacing = line_spacing_rule


def set_paragraph_alignment_properties(paragraph, alignment: str = None): # alignment: "LEFT", "CENTER", "RIGHT", "JUSTIFY"
    """Applies alignment to a paragraph."""
    if alignment:
        try:
            align_enum = getattr(WD_ALIGN_PARAGRAPH, alignment.upper(), None)
            if align_enum is not None:
                paragraph.alignment = align_enum
            else:
                print(f"Warning: Invalid alignment value '{alignment}'. Skipping.")
        except Exception as e:
            print(f"Warning: Exception setting alignment '{alignment}': {e}")

# --- Table Utilities ---

def add_table(doc: Document, rows: int, cols: int, data: list[list[str]] = None, style: str = None) -> Table | None:
    """
    Adds a new table to the document.

    Args:
        doc: The python-docx Document object.
        rows: Number of rows for the table.
        cols: Number of columns for the table.
        data: Optional list of lists of strings to populate the table cells.
              If data dimensions don't match rows/cols, it fills what it can.
        style: Optional name of the table style to apply (e.g., 'Light Shading Accent 1', 'Table Grid').

    Returns:
        The created python-docx.table.Table object, or None if creation fails.
    """
    # Robust type check for Document object due to potential import issues in test environments
    if not (hasattr(doc, 'add_table') and type(doc).__name__ == 'Document' and doc.__class__.__module__.startswith('docx.document')):
        print(f"Error: Invalid Document object provided to add_table. Type was {type(doc).__name__} from module {doc.__class__.__module__}.")
        return None
    if not (isinstance(rows, int) and rows > 0 and isinstance(cols, int) and cols > 0):
        print(f"Error: Invalid rows ({rows}) or cols ({cols}) for add_table. Must be positive integers.")
        return None

    print(f"Attempting to add table with {rows} rows, {cols} cols. Style: {style}")
    try:
        table = doc.add_table(rows=rows, cols=cols)

        if style:
            try:
                table.style = style
                print(f"Applied table style: {style}")
            except Exception as e: # python-docx might raise if style is truly bad, or just apply default
                print(f"Warning: Could not apply table style '{style}'. It might be invalid or not available. Error: {e}")

        if data:
            print("Populating table with provided data.")
            for i in range(rows):
                if i < len(data): # If there's data for this row
                    row_data = data[i]
                    for j in range(cols):
                        if j < len(row_data): # If there's data for this cell in the current row_data
                            table.cell(i, j).text = str(row_data[j]) # Ensure data is string
                        # else: cell remains empty if data[i] is shorter than cols
                # else: row remains empty if data is shorter than rows

        print(f"Table added successfully.")
        return table
    except Exception as e:
        print(f"Error creating table in add_table: {e}")
        return None

def _set_cell_shading(cell: _Cell, hex_color: str):
    """
    Applies background shading to a table cell.
    Requires direct XML manipulation as python-docx does not have a direct API for this.
    """
    if not isinstance(cell, _Cell):
        print("Warning: Invalid cell object provided to _set_cell_shading.")
        return

    # Get the table cell properties element (tcPr)
    tcPr = cell._tc.get_or_add_tcPr()

    # Create a new w:shd element
    shd = OxmlElement('w:shd')

    # Set attributes for the shading
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear') # 'clear' is a common value for solid fills

    # Append the shading element to the cell properties
    tcPr.append(shd)

def format_table_cell(
    table: Table,
    row: int,
    col: int,
    text: str = None,
    font_name: str = None,
    font_size: float = None,
    bold: bool = None,
    italic: bool = None,
    underline: bool = None,
    alignment: str = None,
    shading: str = None
) -> bool:
    """
    Formats a specific cell in a table.

    Args:
        table: The python-docx Table object.
        row: The row index of the cell.
        col: The column index of the cell.
        text: Optional new text for the cell.
        font_name: Optional font name for the cell text.
        font_size: Optional font size (in points) for the cell text.
        bold: Optional bold setting for the cell text.
        italic: Optional italic setting for the cell text.
        underline: Optional underline setting for the cell text.
        alignment: Optional alignment for the cell content ('LEFT', 'CENTER', 'RIGHT', 'JUSTIFY').
        shading: Optional hex color string for the cell's background shading (e.g., 'C0C0C0').

    Returns:
        True if formatting was successful, False otherwise.
    """
    try:
        if row >= len(table.rows) or col >= len(table.columns) or row < 0 or col < 0:
            print(f"Error: Cell ({row}, {col}) is out of bounds for the table.")
            return False

        cell = table.cell(row, col)

        # A cell must have at least one paragraph.
        # We will operate on the first paragraph for text, font, and alignment settings.
        paragraph = cell.paragraphs[0]

        if text is not None:
            # When setting text, clear existing content in the paragraph first.
            # A simple way is to clear runs and add a new one.
            paragraph.text = text
            # After setting the text, we need to re-fetch the paragraph to apply font settings to the new run(s).
            # This is a nuance of python-docx: setting paragraph.text replaces the runs.
            paragraph = cell.paragraphs[0]

        # Apply font properties if any are provided
        if any([font_name, font_size, bold, italic, underline]):
            set_paragraph_font_properties(paragraph, font_name, font_size, bold, italic, underline)

        # Apply alignment if provided
        if alignment:
            set_paragraph_alignment_properties(paragraph, alignment)

        # Apply shading if provided
        if shading:
            _set_cell_shading(cell, shading)

        return True

    except Exception as e:
        print(f"Error formatting cell ({row}, {col}): {e}")
        return False

def merge_table_cells(table: Table, start_row: int, start_col: int, end_row: int, end_col: int) -> bool:
    """
    Merges a rectangular region of cells in a table.

    Args:
        table: The python-docx Table object.
        start_row: The starting row index of the merge region.
        start_col: The starting column index of the merge region.
        end_row: The ending row index of the merge region.
        end_col: The ending column index of the merge region.

    Returns:
        True if the merge was successful, False otherwise.
    """
    try:
        # Basic bounds checking
        if not (0 <= start_row <= end_row < len(table.rows) and 0 <= start_col <= end_col < len(table.columns)):
            print(f"Error: Merge region ({start_row},{start_col}) to ({end_row},{end_col}) is out of bounds.")
            return False

        start_cell = table.cell(start_row, start_col)
        end_cell = table.cell(end_row, end_col)

        start_cell.merge(end_cell)
        print(f"Successfully merged cells from ({start_row},{start_col}) to ({end_row},{end_col}).")
        return True

    except Exception as e:
        # python-docx can raise ValueError for invalid merges (e.g., if region isn't rectangular)
        print(f"Error merging cells: {e}")
        return False

# --- Helper for Scope Resolution ---

def _get_target_paragraphs(doc: Document, elements_details: list, scope: str) -> list:
    """
    Helper function to resolve a scope string to a list of paragraph objects.
    """
    target_paras = []
    if not scope:
        print("Warning: No scope provided for target paragraph resolution.")
        return target_paras

    if scope == "all_paragraphs":
        target_paras = doc.paragraphs
    elif scope.startswith("headings_level_"):
        try:
            level = int(scope.split("_")[-1])
            for el_detail in elements_details:
                if el_detail.get("type") == "heading" and el_detail.get("level") == level:
                    if el_detail["paragraph_index"] < len(doc.paragraphs):
                        target_paras.append(doc.paragraphs[el_detail["paragraph_index"]])
        except ValueError:
            print(f"Warning: Invalid heading level in scope '{scope}'.")
    elif scope.startswith("paragraph_index_"):
        try:
            idx = int(scope.split("_")[-1])
            if 0 <= idx < len(doc.paragraphs):
                target_paras.append(doc.paragraphs[idx])
            else:
                print(f"Warning: Paragraph index {idx} out of bounds for scope '{scope}'.")
        except ValueError:
            print(f"Warning: Invalid paragraph index in scope '{scope}'.")
    elif scope == "all_body_paragraphs":
         for el_detail in elements_details:
            if el_detail.get("type") == "paragraph": # Not a heading
                if el_detail["paragraph_index"] < len(doc.paragraphs):
                    target_paras.append(doc.paragraphs[el_detail["paragraph_index"]])
    else:
        print(f"Warning: Unknown or unsupported scope '{scope}' for paragraph resolution.")

    return target_paras

# --- Action Handlers (Tool Implementations) ---

def apply_find_and_replace_action(doc: Document, action: dict): # Removed unused elements_details
    """
    Finds all occurrences of a word/phrase and replaces it with the given replacement.
    Action: {"action": "find_and_replace", "find": "text_to_find", "replace_with": "replacement_text"}
    """
    print(f"Applying find_and_replace action: {action}")
    find_text = action.get("find")
    replace_with = action.get("replace_with")

    if not find_text or replace_with is None: # Check if find_text is empty or replace_with is not provided (None is a valid replacement)
        print("Warning: 'find' text is empty or 'replace_with' not provided for find_and_replace. Skipping.")
        return

    replaced_count = 0
    for para in doc.paragraphs:
        for run in para.runs:
            if find_text.lower() in run.text.lower():
                pattern = re.compile(re.escape(find_text), re.IGNORECASE)
                new_text, n = pattern.subn(replace_with, run.text)
                if n > 0:
                    run.text = new_text
                    replaced_count += n
    print(f"Applied find_and_replace to {replaced_count} occurrences of '{find_text}'.")

def apply_set_font_action(doc: Document, elements_details: list, action: dict):
    """
    Applies font settings based on the action dictionary.
    Action: {"action": "set_font", "scope": "all_paragraphs" | "headings_level_X" | "paragraph_index_N",
             "font_name": "Arial", "size": 12, "bold": false, "italic": false, "underline": false}
    """
    print(f"Applying font action: {action}")
    scope = action.get("scope")
    font_name = action.get("font_name")
    size = action.get("size")
    bold = action.get("bold")
    italic = action.get("italic")
    underline = action.get("underline")

    target_paras = _get_target_paragraphs(doc, elements_details, scope)
    if not target_paras:
        return # Warning already printed by _get_target_paragraphs

    for para in target_paras:
        set_paragraph_font_properties(para, font_name, size, bold, italic, underline)
    print(f"Applied font to {len(target_paras)} paragraphs for scope '{scope}'.")


def apply_set_heading_style_action(doc: Document, elements_details: list, action: dict):
    """
    Applies style (font, size, bold etc.) to headings of a specific level.
    Action: {"action": "set_heading_style", "level": (int), "font_name": "Calibri Light",
             "size": 18, "bold": true, "spacing_after": 12, "italic": false, "underline": false}
    """
    print(f"Applying heading style action: {action}")
    level = action.get("level")
    font_name = action.get("font_name")
    size = action.get("size")
    bold = action.get("bold")
    italic = action.get("italic")
    underline = action.get("underline")
    spacing_after_pt = action.get("spacing_after")
    # keep_with_next = action.get("keep_with_next") # TODO: Implement if needed

    # Construct scope string for heading level
    scope = f"headings_level_{level}"
    target_paras = _get_target_paragraphs(doc, elements_details, scope)
    if not target_paras:
        return

    for para in target_paras:
        set_paragraph_font_properties(para, font_name, size, bold, italic, underline)
        if spacing_after_pt is not None:
             set_paragraph_spacing_properties(para, spacing_after_pt=spacing_after_pt)
        # if keep_with_next is not None:
        #     para.paragraph_format.keep_with_next = keep_with_next # Requires python-docx feature
    print(f"Applied style to {len(target_paras)} Level {level} headings.")


def apply_set_paragraph_spacing_action(doc: Document, elements_details: list, action: dict):
    """
    Applies paragraph spacing settings.
    Action: {"action": "set_paragraph_spacing", "scope": "all_paragraphs",
             "spacing_before": 0 (pt), "spacing_after": 6 (pt), "line_spacing": 1.15 (rule)}
    """
    print(f"Applying paragraph spacing action: {action}")
    scope = action.get("scope")
    spacing_before = action.get("spacing_before")
    spacing_after = action.get("spacing_after")
    line_spacing = action.get("line_spacing")

    target_paras = _get_target_paragraphs(doc, elements_details, scope)
    if not target_paras:
        return

    for para in target_paras:
        set_paragraph_spacing_properties(para, spacing_before, spacing_after, line_spacing)
    print(f"Applied spacing to {len(target_paras)} paragraphs for scope '{scope}'.")


def apply_set_alignment_action(doc: Document, elements_details: list, action: dict):
    """
    Applies text alignment.
    Action: {"action": "set_alignment", "scope": "headings_level_1", "alignment": "LEFT" | "CENTER" | "RIGHT" | "JUSTIFY"}
    """
    print(f"Applying alignment action: {action}")
    scope = action.get("scope")
    alignment = action.get("alignment")

    target_paras = _get_target_paragraphs(doc, elements_details, scope)
    if not target_paras:
        return

    for para in target_paras:
        set_paragraph_alignment_properties(para, alignment)
    print(f"Applied alignment to {len(target_paras)} paragraphs for scope '{scope}'.")

# TODO: Implement other action handlers like:
# def apply_ensure_consistent_style_action(doc: Document, elements_details: list, action: dict): ...
# def apply_theme_action(doc: Document, elements_details: list, action: dict): ...

def apply_fix_font_inconsistencies_action(doc: Document, elements_details: list, action: dict): # elements_details might not be needed if we iterate all paras
    """
    Attempts to unify fonts across the document based on a target font and size.
    Action: {"action": "fix_font_inconsistencies", "target_font_name": "Calibri", "target_font_size": 11}
    This is a broad approach. More nuanced logic could be added to preserve specific formatting
    (e.g., for code blocks, or manually emphasized text if distinguishable).
    """
    print(f"Applying font inconsistency fix: {action}")
    target_font_name = action.get("target_font_name")
    target_font_size_pt = action.get("target_font_size") # Assuming this is in Pt

    if not target_font_name and not target_font_size_pt:
        print("Warning: No target font name or size provided for fix_font_inconsistencies. Skipping.")
        return

    changed_elements_count = 0
    # This action, by its nature, typically applies document-wide or to all body text.
    # Using _get_target_paragraphs can allow for more specific scoping if the LLM/plan provides it.
    # If no scope is given, or "all_paragraphs", it will iterate through all.
    scope = action.get("scope", "all_paragraphs") # Default to all_paragraphs if no scope provided

    target_paras = _get_target_paragraphs(doc, elements_details, scope)
    if not target_paras: # If scope was invalid or no paras matched
        if not scope or scope == "all_paragraphs": # If it intended to run on all but list is empty
             print("Warning: No paragraphs found to apply font inconsistency fix, or document is empty.")
        return

    for paragraph in target_paras:
        # The original logic iterated elements_details and then used doc.paragraphs[para_idx].
        # Iterating target_paras (which are actual paragraph objects) is more direct.
        for run in paragraph.runs:
            applied_change_to_run = False
            if target_font_name and run.font.name != target_font_name:
                run.font.name = target_font_name
                applied_change_to_run = True
            if target_font_size_pt and (not run.font.size or run.font.size.pt != target_font_size_pt): # Check if size exists before comparing
                run.font.size = Pt(target_font_size_pt)
                applied_change_to_run = True
            if applied_change_to_run:
                changed_elements_count +=1 # Count runs changed

    print(f"Applied font inconsistency fix to {changed_elements_count} runs within scope '{scope}'.")


if __name__ == '__main__':
    # This is for basic testing of the utility functions.
    # You would need a sample .docx file named 'sample.docx' in the same directory.
    # Create a dummy document for testing if 'sample.docx' doesn't exist
    doc_main = Document() # Renamed to avoid conflict with 'doc' in add_table test
    doc_main.add_heading('Test Document', level=0)
    doc_main.add_paragraph('This is a test paragraph.')
    doc_main.add_heading('Heading 1', level=1)
    doc_main.add_paragraph('Another paragraph under Heading 1.')
    doc_main.add_heading('Heading 2', level=2)
    doc_main.add_paragraph('Paragraph under Heading 2.')
    save_document(doc_main, 'sample_test_doc.docx')

    print("Testing with 'sample_test_doc.docx'")
    loaded_doc = load_document('sample_test_doc.docx')

    print("\nExtracted Paragraph Texts:")
    texts = extract_text_from_paragraphs(loaded_doc)
    for text in texts:
        print(f"- {text}")

    print("\nExtracted Headings:")
    headings_info = extract_headings(loaded_doc)
    for heading in headings_info:
        print(f"- Level {heading['level']}: {heading['text']}")

    print("\nDetailed Document Analysis:")
    detailed_analysis_data = get_document_analysis(loaded_doc) # Use the loaded_doc directly
    for item in detailed_analysis_data["elements"]:
        print(f"- Type: {item['type']}")
        if item['type'] == 'heading':
            print(f"  Level: {item['level']}")
        print(f"  Text: {item['text'][:50]}...") # Print first 50 chars
        if 'style_name' in item: # Corrected from 'style' to 'style_name' for consistency
            print(f"  Style: {item['style_name']}")
        if 'runs' in item:
            if item['runs']:
                first_run = item['runs'][0]
                print(f"  First run font: {first_run.get('font_name', 'N/A')}, Size: {first_run.get('font_size', 'N/A')}, Bold: {first_run.get('bold', 'N/A')}")

    # Test add_table
    print("\n--- Testing add_table utility ---")
    doc_for_table_test = Document()
    doc_for_table_test.add_paragraph("Paragraph before table.")
    table_data = [["Name", "Role"], ["Jules", "Agent"], ["User", "Supervisor"]]
    created_table = add_table(doc_for_table_test, rows=3, cols=2, data=table_data, style="Table Grid")

    if created_table:
        print(f"Test table created with style: {created_table.style.name if created_table.style else 'Default'}")
        # Basic check of table content
        if hasattr(created_table, 'rows') and hasattr(created_table, 'columns') and \
           len(created_table.rows) == 3 and len(created_table.columns) == 2:
            print(f"  Cell (0,0): {created_table.cell(0,0).text}")
            print(f"  Cell (1,1): {created_table.cell(1,1).text}")
        else:
            print("  Error: Table dimensions mismatch or table object incomplete after creation.")
    else:
        print("Test table creation FAILED.")
    doc_for_table_test.add_paragraph("Paragraph after table.")

    save_document(doc_for_table_test, 'sample_test_doc_with_table.docx')
    print("Saved 'sample_test_doc_with_table.docx' for manual inspection.")

    # Clean up the dummy files
    import os
    if os.path.exists('sample_test_doc.docx'):
        os.remove('sample_test_doc.docx')
        print("\nCleaned up sample_test_doc.docx")
    if os.path.exists('sample_test_doc_with_table.docx'):
        os.remove('sample_test_doc_with_table.docx')
        print("Cleaned up sample_test_doc_with_table.docx")
